from io import BytesIO

from pandas import DataFrame

import spacy
from spacy.tokens import SpanGroup
from docx import Document
from docx.table import Table
from clause_segmenter import ClauseSegmenter

CLAUSE_FIELD: str = "clause"
START_FIELD: str = "start_idx"
END_FIELD: str = "end_idx"


class SourceFileClauser:
    NLP = spacy.load("en_core_web_sm")

    def __init__(self, source_file: BytesIO, filetype: str):
        self.source_text: str
        if filetype == "docx":
            self.source_text = SourceFileClauser.read_docx(source_file)
        elif filetype == "txt":
            self.source_text = SourceFileClauser.read_txt(source_file)
        else:
            raise ValueError(f"File type {filetype} is not a valid source file type")
        self.clause_segmenter = ClauseSegmenter(self.NLP)
        self.doc = None

    def get_text(self) -> str:
        return self.source_text

    @staticmethod
    def read_docx(docx_file: BytesIO) -> str:
        docx_file.seek(0)
        try:
            doc = Document(docx_file)
        except Exception:
            raise ValueError("Loaded file type or file format is incorrect")
        doc_text: str = ""

        if len(doc.tables) == 0:
            raise ValueError("No table found within the loaded document")

        text_table: Table = doc.tables[0]

        content_col: int = -1
        for idx, cell in enumerate(text_table.rows[0].cells):
            if cell.text.lower() == "content":
                content_col = idx

        if content_col == -1:
            raise ValueError("No column labelled 'content' found in table")

        for row in text_table.rows[1:]:
            cell = row.cells[content_col]
            doc_text += f"{cell.text} \n"

        return doc_text

    @staticmethod
    def read_txt(txt_file: BytesIO) -> str:
        txt_file.seek(0)
        txt_bytes: bytes = txt_file.read()
        try:
            txt_content: str = txt_bytes.decode('utf-8')
        except UnicodeDecodeError:
            raise ValueError("Text file must be encoded in UTF-8")

        return txt_content

    def generate_clause_dataframe(self) -> DataFrame:
        if self.doc is None:
            self.doc = SourceFileClauser.NLP(self.source_text)
        clauses: SpanGroup = self.clause_segmenter.get_clauses_as_spangroup(self.doc)
        clause_data: dict[str, list] = {CLAUSE_FIELD: [], START_FIELD: [], END_FIELD: []}
        for clause_span in clauses:
            clause_data[CLAUSE_FIELD].append(clause_span.text_with_ws)
            clause_data[START_FIELD].append(clause_span.start_char)
            clause_data[END_FIELD].append(clause_span.end_char)

        return DataFrame(clause_data, columns=[CLAUSE_FIELD, START_FIELD, END_FIELD])
